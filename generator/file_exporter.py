# generator/file_exporter.py
import io
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageDraw
from moviepy.editor import TextClip

def generate_pdf(prompt: str, content: str) -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    c.drawString(100, height - 100, prompt)
    c.drawString(100, height - 120, content)
    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer.getvalue()

def generate_word(prompt: str, content: str) -> bytes:
    document = Document()
    document.add_heading(prompt, level=1)
    document.add_paragraph(content)
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generate_excel(prompt: str, data: list) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = prompt
    for row in data:
        ws.append(row)
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generate_image(prompt: str, content: str) -> bytes:
    img = Image.new('RGB', (400, 300), color = (73, 109, 137))
    draw = ImageDraw.Draw(img)
    draw.text((10, 10), prompt, fill=(255,255,0))
    buffer = io.BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer.getvalue()

def generate_video(prompt: str, content: str, duration=5) -> bytes:
    clip = TextClip(f"{prompt}\n{content}", fontsize=24, color='white', bg_color='black', size=(640,480))
    clip = clip.set_duration(duration)
    temp_filename = "temp_video.mp4"
    clip.write_videofile(temp_filename, codec="libx264", audio=False, verbose=False, logger=None)
    with open(temp_filename, "rb") as f:
        video_data = f.read()
    import os
    os.remove(temp_filename)
    return video_data
